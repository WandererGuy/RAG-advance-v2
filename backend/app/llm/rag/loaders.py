"""File -> pages. The only place that knows about PDF or DOCX internals.

Loaders return text exactly as the file yields it. No header/footer stripping, no de-hyphenation,
no whitespace surgery beyond stripping trailing spaces on a line: the v1 corpus was inspected
page by page and carries no repeated running header or footer, so a stripping heuristic here
would only be a way to delete real content by accident. Add one when a document needs it, and
write an ADR when you do.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from app.core.exceptions import UnsupportedFileType

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass(frozen=True)
class Page:
    """One unit of text with the page number a citation will point at.

    `page_no` is 1-based to match what a reader sees in a PDF viewer. It is None only for
    formats that have no real pagination — see `load_docx`.
    """

    page_no: int | None
    text: str


@dataclass(frozen=True)
class LoadedDocument:
    pages: list[Page]
    mime_type: str
    # None where the format has no page concept; `documents.page_count` is nullable for this.
    page_count: int | None

    @property
    def total_chars(self) -> int:
        return sum(len(page.text) for page in self.pages)


def _normalize(text: str) -> str:
    """Strip trailing whitespace per line and drop leading/trailing blank lines.

    Deliberately minimal. Interior line breaks are left alone: PDF text is hard-wrapped
    mid-sentence, and joining those lines would be a guess about where a paragraph really ends.
    The chunker treats a line break as a usable split boundary, so the wrapping costs nothing.
    """
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


def load_pdf(path: Path) -> LoadedDocument:
    """Load a text-based PDF with PyMuPDF, preserving real page numbers.

    Scanned PDFs are out of scope for v1 (no OCR): they load as pages of empty text rather
    than raising, and the ingest service is what rejects a document with no extractable text.
    """
    import pymupdf

    pages: list[Page] = []
    with pymupdf.open(path) as document:
        for index in range(document.page_count):
            pages.append(Page(page_no=index + 1, text=_normalize(document[index].get_text())))

    return LoadedDocument(pages=pages, mime_type=PDF_MIME, page_count=len(pages))


def load_docx(path: Path) -> LoadedDocument:
    """Load a DOCX with python-docx.

    **A DOCX has no page numbers.** Pagination in Word is computed by the renderer from fonts,
    margins and printer metrics; the file itself stores only a stream of paragraphs. Everything
    this loader returns therefore carries `page_no=None`, and citations from a DOCX name the
    document without a page. `chunks.page_no` is nullable for exactly this reason.

    The v1 corpus is PDF-only, so this path is not exercised by the corpus.
    """
    import docx

    paragraphs = [p.text for p in docx.Document(str(path)).paragraphs]
    text = _normalize("\n".join(paragraphs))
    return LoadedDocument(
        pages=[Page(page_no=None, text=text)], mime_type=DOCX_MIME, page_count=None
    )


_LOADERS = {".pdf": load_pdf, ".docx": load_docx}

SUPPORTED_SUFFIXES = frozenset(_LOADERS)


def load(path: Path) -> LoadedDocument:
    """Dispatch on file extension. Raises UnsupportedFileType for anything else."""
    suffix = path.suffix.lower()
    loader = _LOADERS.get(suffix)
    if loader is None:
        raise UnsupportedFileType(str(path), suffix)
    return loader(path)
